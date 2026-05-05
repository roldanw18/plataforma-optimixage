"""
Genera el informe de Sprint 2 — versión resumida con normas APA 7.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── PÁGINA ─────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Cm(21.59)
sec.page_height   = Cm(27.94)
sec.top_margin    = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin   = Cm(2.54)
sec.right_margin  = Cm(2.54)

# ── ESTILO BASE ─────────────────────────────────────────────────
def norm():
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(0)
norm()

TNR = 'Times New Roman'

def _p(align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, li=None, fi=None, ls=24):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before       = Pt(before)
    pf.space_after        = Pt(after)
    pf.line_spacing       = Pt(ls)
    if li is not None: pf.left_indent            = Cm(li)
    if fi is not None: pf.first_line_indent      = Cm(fi)
    p.alignment = align
    return p

def _r(p, text, bold=False, italic=False, size=12):
    r = p.add_run(text)
    r.font.name  = TNR
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    return r

def h1(text):
    p = _p(WD_ALIGN_PARAGRAPH.CENTER, before=18, after=6)
    _r(p, text.upper(), bold=True)

def h2(text):
    p = _p(before=12, after=4)
    _r(p, text, bold=True)

def h3(text):
    p = _p(before=8, after=3)
    _r(p, text, bold=True, italic=True)

def body(text):
    p = _p(fi=1.27)
    _r(p, text)

def blank():
    _p()

def blt(text, lvl=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(1.27 + lvl*0.63)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(24)
    _r(p, text)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.name = TNR; r.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = TNR; r.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    blank()

# ════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════
for _ in range(3): blank()
p = _p(WD_ALIGN_PARAGRAPH.CENTER); _r(p, "Informe de Avance — Sprint 2", bold=True, size=14)
blank()
p = _p(WD_ALIGN_PARAGRAPH.CENTER); _r(p, "Plataforma de Seguimiento de Proyectos", bold=True, size=13)
blank(); blank()
for line in ["Wbeimar", "", "Ingeniería de Software — Práctica Profesional", "", "15 de abril de 2026"]:
    p = _p(WD_ALIGN_PARAGRAPH.CENTER); _r(p, line)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# RESUMEN
# ════════════════════════════════════════════════════════
h1("Resumen")
blank()
body(
    "Este informe resume los resultados del Sprint 2 de la Plataforma de Seguimiento de "
    "Proyectos, una aplicación web para gestionar el avance de proyectos de digitalización "
    "entre administradores y clientes. Durante este sprint se construyó e integró el "
    "frontend en React —tomando como referencia los diseños que la empresa entregó al inicio "
    "de la práctica y usando v0.dev de Vercel para generar el código base de las interfaces—, "
    "se ampliaron las funcionalidades del backend y se consolidó la infraestructura de CI/CD. "
    "El sprint completó todos sus ítems del Product Backlog y adelantó funcionalidades "
    "previstas para el Sprint 3."
)
blank()
p = _p(fi=0)
_r(p, "Palabras clave: ", bold=True, italic=True)
_r(p, "gestión de proyectos, React, v0.dev, FastAPI, CI/CD, control de acceso, notificaciones", italic=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ════════════════════════════════════════════════════════
h1("1. Introducción")
blank()
body(
    "La Plataforma de Seguimiento de Proyectos nació para centralizar toda la información "
    "de un proyecto —etapas, documentos, tareas, reuniones y notificaciones— en un solo "
    "lugar, reemplazando el uso disperso de correos y aplicaciones de mensajería. El Sprint 1 "
    "construyó los cimientos: autenticación JWT, control de acceso por roles y pruebas "
    "automáticas. El Sprint 2, documentado aquí, convirtió ese prototipo técnico en una "
    "aplicación usable por personas reales, con pantallas para clientes y administradores "
    "conectadas a todos los módulos del backend."
)

# ════════════════════════════════════════════════════════
# 2. OBJETIVOS
# ════════════════════════════════════════════════════════
h1("2. Objetivos del Sprint 2")
blank()
h2("2.1. Objetivo General")
body(
    "Conectar el frontend de la plataforma —construido en React con apoyo de v0.dev y "
    "los diseños iniciales entregados por la empresa— con el backend FastAPI, e "
    "implementar los módulos funcionales definidos en el Product Backlog para el Sprint 2."
)
blank()
h2("2.2. Ítems del Product Backlog Comprometidos (Sprint 2)")
table(
    ["ID", "Nombre", "Tipo", "Estado"],
    [
        ["US03", "Gestión de Proyectos (Crear y Editar)",      "Feature",        "Completado"],
        ["US04", "Visualización de Proyecto Asignado",          "Feature",        "Completado"],
        ["US05", "Gestión de Proyectos Activos",                "Feature",        "Completado"],
        ["US06", "Visualización del Proceso y Etapa Actual",    "Feature",        "Completado"],
        ["US07", "Cambio de Estado de Etapa",                   "Feature",        "Completado"],
        ["US08", "Historial de Fechas por Etapa",               "Feature",        "Completado"],
        ["TW05", "Validación y Seguridad de Entradas",          "Technical Work", "Completado"],
        ["TW06", "Manejo Centralizado de Errores",              "Technical Work", "Completado"],
    ],
    widths=[1.8, 7.5, 3.2, 3.0]
)
h2("2.3. Ítems Adelantados del Sprint 3")
body(
    "Además de los comprometidos, el equipo adelantó trabajo de varios ítems del Sprint 3, "
    "lo que reduce la carga del siguiente sprint."
)
blank()
table(
    ["ID", "Nombre", "Estado en Sprint 2"],
    [
        ["US09", "Gestión y Visualización de Documentos", "Base implementada (falta carga real de archivos)"],
        ["US11", "Módulo de Comunicación",                "Mensajería básica implementada"],
        ["US12", "Dashboard Informativo Ampliado",        "Completado"],
        ["US14", "Gestión de Contenido Formativo",        "Estructura base (sin reproductor de video)"],
        ["TW09", "Logging y Auditoría del Sistema",       "Completado (AuditLog + EtapaHistorial)"],
    ],
    widths=[1.8, 6.5, 7.2]
)

# ════════════════════════════════════════════════════════
# 3. DISEÑO DEL FRONTEND
# ════════════════════════════════════════════════════════
h1("3. Diseño del Frontend")
blank()
body(
    "El diseño visual de las interfaces partió de los mockups y materiales de referencia que "
    "la empresa entregó al comienzo de la práctica. Con esa base, se usó v0.dev de Vercel "
    "para generar el código inicial de los componentes React, lo que aceleró "
    "significativamente el desarrollo. Después, ese código fue adaptado, ajustado e "
    "integrado al proyecto para que funcionara con la API real y el sistema de autenticación."
)
blank()
body(
    "Este flujo de trabajo —diseños iniciales → v0.dev → ajuste e integración— permitió "
    "tener interfaces con buena apariencia desde el principio, sin dedicar tiempo excesivo "
    "a construir componentes desde cero."
)

# ════════════════════════════════════════════════════════
# 4. STACK TECNOLÓGICO
# ════════════════════════════════════════════════════════
h1("4. Stack Tecnológico")
blank()
table(
    ["Área", "Tecnología", "Detalle"],
    [
        ["Backend / API",      "Python + FastAPI",        "Python 3.10, FastAPI ≥ 0.100"],
        ["ORM / Migraciones",  "SQLAlchemy + Alembic",    "SQLAlchemy 2.0"],
        ["Base de datos",      "PostgreSQL / SQLite",     "PostgreSQL en prod · SQLite en tests"],
        ["Autenticación",      "JWT + bcrypt",            "python-jose · passlib"],
        ["Frontend",           "React + Vite",            "React 19, Vite 6"],
        ["Generación de UI",   "v0.dev (Vercel)",         "Código base de componentes React"],
        ["Estilos",            "Tailwind CSS",            "Tailwind CSS 3"],
        ["Comunicación API",   "Axios",                   "Interceptores JWT, manejo de 401"],
        ["Testing",            "Pytest + pytest-cov",     "Cobertura ≥ 60 %"],
        ["Contenedores",       "Docker + Docker Compose", "3 servicios: DB, Backend, Frontend"],
        ["CI/CD",              "GitHub Actions",          "Push a main/develop/feature/*"],
    ],
    widths=[4.2, 4.2, 7.1]
)

# ════════════════════════════════════════════════════════
# 5. FUNCIONALIDADES IMPLEMENTADAS
# ════════════════════════════════════════════════════════
h1("5. Funcionalidades Implementadas")
blank()

h2("5.1. Backend — Nuevos Módulos")
blank()
table(
    ["Módulo", "Qué hace", "Quién puede usarlo"],
    [
        ["Proceso",        "Mueve el proyecto entre 5 etapas (20 % → 100 %) y guarda el historial de cambios.", "Admin cambia · Cliente consulta"],
        ["Documentos",     "Crea y publica documentos del proyecto (contratos, informes, propuestas).",          "Admin crea · Cliente consulta"],
        ["Mensajería",     "Canal de mensajes por proyecto con estado de lectura.",                              "Ambos roles"],
        ["Hitos",          "Define y actualiza los compromisos del proyecto (pendiente / en progreso / listo).", "Admin gestiona · Cliente consulta"],
        ["Tareas",         "Asigna tareas con prioridad (baja–urgente) y estado a usuarios del proyecto.",       "Admin crea · Asignado actualiza estado"],
        ["Reuniones",      "Programa y cancela reuniones; genera notificación automática al crearlas.",          "Admin gestiona"],
        ["Notificaciones", "Bandeja de entrada por usuario, disparada automáticamente por eventos del sistema.", "Ambos roles"],
    ],
    widths=[3.2, 8.8, 4.5]
)

h2("5.2. Frontend — Páginas")
blank()
h3("Portal del Cliente")
client = [
    ("Inicio",       "Dashboard con documentos recientes, barra de progreso y próxima reunión."),
    ("Proceso",      "Línea de tiempo de las 5 etapas, porcentaje de avance e historial de cambios."),
    ("Documentos",   "Galería de documentos publicados con íconos por tipo."),
    ("Contenido",    "Sección de materiales formativos (estructura base lista)."),
    ("Contacto",     "Formulario de soporte con el equipo administrador."),
]
for name, desc in client:
    p = _p(li=1.27)
    _r(p, f"{name}. ", bold=True)
    _r(p, desc)

blank()
h3("Panel de Administración")
admin = [
    ("Clientes",        "Lista de clientes, creación de proyectos y asignación cliente–proyecto."),
    ("Proceso (Admin)", "Vista de proyectos activos con modal para cambiar etapa y registrar notas."),
    ("Notificaciones",  "Centro de notificaciones del sistema con marcado individual y masivo."),
    ("Contenido",       "Gestión de materiales de capacitación."),
]
for name, desc in admin:
    p = _p(li=1.27)
    _r(p, f"{name}. ", bold=True)
    _r(p, desc)

blank()
h2("5.3. Base de Datos — Modelos Nuevos")
blank()
table(
    ["Modelo", "Para qué sirve"],
    [
        ["Documento",      "Archivos y su metadata asociados a un proyecto."],
        ["Mensaje",        "Mensajes del proyecto con estado de lectura."],
        ["Hito",           "Compromisos de entrega con estado."],
        ["Tarea",          "Tareas asignables con prioridad y estado."],
        ["Reunion",        "Reuniones programadas."],
        ["Notificacion",   "Avisos en bandeja de cada usuario."],
        ["EtapaHistorial", "Registro de cada cambio de etapa (quién, cuándo, nota)."],
        ["AuditLog",       "Registro de acciones del sistema para auditoría."],
    ],
    widths=[4.5, 12.0]
)

h2("5.4. Pruebas Automatizadas")
blank()
table(
    ["Archivo de prueba", "Qué verifica", "Resultado"],
    [
        ["test_proceso.py",        "Cambio de etapas, permisos, historial, progreso 100 %", "Pasando"],
        ["test_documentos.py",     "Creación y consulta de documentos",                     "Pasando"],
        ["test_mensajes.py",       "Envío y recuperación de mensajes",                      "Pasando"],
        ["test_proyectos.py",      "Operaciones básicas de proyectos",                      "Pasando"],
        ["test_proyectos_admin.py","Operaciones de administrador",                           "Pasando"],
    ],
    widths=[5.0, 8.5, 3.0]
)

# ════════════════════════════════════════════════════════
# 6. ESTADO DEL SPRINT
# ════════════════════════════════════════════════════════
h1("6. Estado del Sprint 2")
blank()
body(
    "Todos los ítems comprometidos en el Product Backlog para el Sprint 2 fueron "
    "completados. Adicionalmente, se adelantó parte del trabajo del Sprint 3, lo que "
    "deja al equipo en una posición ventajosa para el siguiente ciclo."
)
blank()
table(
    ["Área", "Estado"],
    [
        ["8 ítems del backlog de Sprint 2 (US03–US08, TW05, TW06)",   "Completados"],
        ["5 ítems del backlog de Sprint 3 (parcial/base)",             "Adelantados"],
        ["Frontend: Portal cliente (5 páginas)",                       "Completo"],
        ["Frontend: Panel admin (4 páginas)",                          "Completo"],
        ["Backend: 7 nuevos módulos",                                  "Completo"],
        ["Base de datos: 8 nuevos modelos + migraciones",              "Completo"],
        ["Pruebas automáticas (5 módulos)",                            "Completo"],
        ["CI/CD GitHub Actions",                                       "Completo"],
        ["Docker Compose (3 servicios)",                               "Completo"],
        ["Carga real de archivos (upload binario)",                    "Pendiente Sprint 3"],
        ["Reproductor de video con proveedor real",                    "Pendiente Sprint 3"],
        ["Notificaciones en tiempo real (WebSockets)",                 "Pendiente Sprint 3"],
    ],
    widths=[11.5, 5.0]
)

# ════════════════════════════════════════════════════════
# 7. COSAS QUE QUEDARON PENDIENTES
# ════════════════════════════════════════════════════════
h1("7. Cosas Que Quedaron Pendientes")
blank()
body(
    "Hay tres puntos que no se terminaron en este sprint y que vale la pena explicar:"
)
blank()
h3("Carga real de archivos")
body(
    "Los documentos se guardan con una URL de referencia en la base de datos, pero todavía "
    "no hay una manera de subir archivos desde la plataforma. Para el Sprint 3 se necesita "
    "conectar un servicio de almacenamiento como Amazon S3, Azure Blob Storage o un "
    "sistema de archivos local con Nginx."
)
blank()
h3("Reproductor de video")
body(
    "La sección de contenido formativo ya tiene la estructura y los botones de reproducción, "
    "pero no está conectada a ningún proveedor de video. En el Sprint 3 se integrará con "
    "Vimeo, YouTube privado o almacenamiento en la nube."
)
blank()
h3("Notificaciones en tiempo real")
body(
    "Las notificaciones funcionan bien, pero el sistema hace consultas periódicas al "
    "servidor para actualizarse. Lo ideal es que lleguen solas en tiempo real usando "
    "WebSockets, lo que está planeado para el Sprint 3."
)

# ════════════════════════════════════════════════════════
# 8. CÓMO QUEDA EL SISTEMA PARA EL SPRINT 3
# ════════════════════════════════════════════════════════
h1("8. Cómo Queda el Sistema para el Sprint 3")
blank()
body(
    "El Sprint 2 deja la plataforma en muy buen estado para seguir construyendo. "
    "Lo que hay disponible y funcionando:"
)
blank()
for item in [
    "API REST con más de 10 módulos, con documentación automática en /docs.",
    "Base de datos con 12 tablas, relaciones e historial de auditoría.",
    "Frontend con componentes reutilizables y sesión persistente.",
    "Sistema de notificaciones que se puede extender con nuevos tipos de eventos.",
    "Pipeline CI/CD estable que detecta errores antes de llegar a producción.",
    "Docker Compose listo para desplegar en cualquier proveedor de nube.",
]:
    blt(item)
blank()

h2("8.1. Backlog del Sprint 3")
blank()
table(
    ["ID", "Nombre", "Prioridad"],
    [
        ["US09",  "Gestión y descarga real de documentos",      "Alta"],
        ["US10",  "Carga de documentos (upload)",               "Alta"],
        ["US11",  "Módulo de comunicación (chat completo)",     "Alta"],
        ["US13",  "Gestión de perfil y cambio de contraseña",   "Media"],
        ["US14",  "Contenido formativo con reproductor real",   "Media"],
        ["TW07",  "Diseño responsive (móvil y tablet)",         "Media"],
        ["TW08",  "Optimización de rendimiento y paginación",   "Media"],
        ["—",     "WebSockets para notificaciones en tiempo real","Alta"],
        ["—",     "Despliegue en la nube (Azure / AWS)",        "Alta"],
    ],
    widths=[2.0, 10.0, 3.5]
)

# ════════════════════════════════════════════════════════
# 9. QUÉ SE PODRÍA MEJORAR EN EL FUTURO
# ════════════════════════════════════════════════════════
h1("9. Qué Se Podría Mejorar en el Futuro")
blank()
h3("A corto plazo (Sprint 3)")
for item in [
    "Subida real de archivos con almacenamiento en la nube.",
    "Notificaciones que llegan solas (WebSockets).",
    "Módulo de presupuestos y cotizaciones.",
    "Despliegue en producción.",
]:
    blt(item)
blank()
h3("A mediano plazo")
for item in [
    "App móvil para seguimiento desde el celular.",
    "Dashboard con métricas y estadísticas de los proyectos.",
    "Integración con herramientas como Jira o Monday.com.",
    "Notificaciones por correo electrónico.",
    "Firma electrónica de documentos.",
]:
    blt(item)
blank()
h3("A largo plazo")
for item in [
    "Inteligencia artificial para detectar riesgos y retrasos.",
    "Modelo SaaS con soporte para múltiples organizaciones.",
    "API pública para integraciones de terceros.",
]:
    blt(item)

# ════════════════════════════════════════════════════════
# 10. CONCLUSIONES
# ════════════════════════════════════════════════════════
h1("10. Conclusiones")
blank()
body(
    "El Sprint 2 cumplió con lo que se comprometió en el Product Backlog y fue más allá, "
    "adelantando trabajo del siguiente sprint. La combinación de los diseños entregados "
    "por la empresa y el uso de v0.dev para generar código base permitió construir "
    "interfaces de calidad en menos tiempo."
)
blank()
body(
    "La arquitectura por capas que se definió desde el inicio demostró ser la decisión "
    "correcta: agregar siete nuevos módulos al backend no requirió tocar la lógica "
    "central, y el frontend se integró sin complicaciones gracias a la claridad de la API."
)
blank()
body(
    "Los puntos que quedaron sin terminar —carga de archivos, video y tiempo real— no "
    "afectan el funcionamiento del sistema y tienen un camino claro de solución para el "
    "Sprint 3. El equipo entra al siguiente sprint con una plataforma sólida, bien "
    "probada y lista para recibir las funcionalidades más avanzadas."
)

# ════════════════════════════════════════════════════════
# REFERENCIAS (APA 7)
# ════════════════════════════════════════════════════════
doc.add_page_break()
h1("Referencias")
blank()

refs = [
    "FastAPI. (2024). FastAPI — framework, alto rendimiento, fácil de aprender, rápido de programar, listo para producción. https://fastapi.tiangolo.com/",
    "Meta Open Source. (2024). React — The library for web and native user interfaces. https://react.dev/",
    "Vercel. (2024). v0.dev — Generative UI. https://v0.dev/",
    "Vite. (2024). Vite — Next generation frontend tooling. https://vitejs.dev/",
    "SQLAlchemy. (2024). SQLAlchemy — The Python SQL toolkit and Object Relational Mapper. https://www.sqlalchemy.org/",
    "Docker Inc. (2024). Docker Compose overview. https://docs.docker.com/compose/",
    "GitHub. (2024). GitHub Actions — Automate your workflow from idea to production. https://github.com/features/actions/",
    "Tailwind Labs. (2024). Tailwind CSS — Rapidly build modern websites without ever leaving your HTML. https://tailwindcss.com/",
    "American Psychological Association. (2020). Publication manual of the American Psychological Association (7.ª ed.). https://doi.org/10.1037/0000165-000",
    "Pressman, R. S., & Maxim, B. R. (2020). Ingeniería del software: Un enfoque práctico (9.ª ed.). McGraw-Hill Education.",
    "Sommerville, I. (2016). Software engineering (10.ª ed.). Pearson Education.",
]
for ref in refs:
    p = _p(li=1.27, fi=-1.27)
    _r(p, ref)

# ════════════════════════════════════════════════════════
# GUARDAR
# ════════════════════════════════════════════════════════
out = r"c:\Users\USUARIO\Desktop\d\docs\Informe_Sprint2_APA_v2.docx"
doc.save(out)
print(f"Listo: {out}")
